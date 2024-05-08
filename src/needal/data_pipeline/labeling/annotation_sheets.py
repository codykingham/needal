"""Module for generating and ingesting annotation sheets."""

import docx
import json
import hashlib

from pathlib import Path
from abc import ABC, abstractmethod
from typing import List, TypedDict, Optional, Union, Dict, Any, Tuple
from tf.fabric import Fabric
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from kingham_thesis.data_pipeline.labeling.specifiers import LingLabel


class AnnotationSheetSpecs(TypedDict):
    """TypedDict to hold specs for an annotation sheet."""
    sheet: str
    project: str


def add_hyperlink(paragraph, url, text, color=None, underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    Source: https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :param color: a color for the URL
    :param underline: whether to underline the URL
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


class BaseAnnotationSheet(ABC):
    """Object for generating an annotation sheet as a MS document."""

    NAME = "base"
    HASH_LEN = 8

    def __init__(
            self,
            annotations: List[LingLabel],
            tf_fabric: Fabric,
            project: 'BaseLabelingProject',
            document: Optional[Document] = None,
    ) -> None:
        """Initialize an AnnotationSheet object."""
        self.annotations = annotations
        self.annotation_metadata, self.label_to_id = (
            self._get_annotation_metadata(annotations)
        )
        self.project = project
        self.tf_fabric = tf_fabric
        self.tf_api = tf_fabric.api
        self.styles = {}
        if not document:
            self.document = Document()
            self._inject_specs_into_docx_metadata(self.document)
            self._add_styles()
            self._build_document(self.document)
        else:
            self.document = document

    @property
    def specs(self) -> AnnotationSheetSpecs:
        """Retrieve specs for this sheet template."""
        return {
            "sheet": self.NAME,
            "project": self.project.name,
        }

    def _inject_specs_into_docx_metadata(self, doc: Document):
        """Add annotation specs to the document's "comments" metadata section."""
        doc.core_properties.comments = json.dumps(self.specs)

    def to_docx(self, filepath: Path) -> None:
        """Save a new docx with document."""
        self.document.save(filepath)

    @abstractmethod
    def _add_styles(self) -> None:
        """Set all styles for the document."""

    @abstractmethod
    def _build_document(self, document: Document) -> None:
        """Build up document."""

    @staticmethod
    @abstractmethod
    def _label_from_row(row, metadata: Dict[str, Any]) -> LingLabel:
        """Extract cell values from a table row into a LingLabel object."""

    def _get_new_hash_id(
            self,
            label: LingLabel,
            annotation_metadata: Dict[str, Dict[str, Any]],
    ) -> str:
        """Get shortened annotation hash ID."""
        # turn label into string and hash it
        label_str = (
            label.label + str(label.nid) + label.target
        )
        label_hash = hashlib.sha1(label_str.encode()).hexdigest()

        # get unique truncated hash to prevent collisions
        trunc_len = self.HASH_LEN
        truncated_hash = label_hash[:trunc_len]
        while truncated_hash in annotation_metadata:
            trunc_len += 1
            truncated_hash = label_hash[:trunc_len]
        return truncated_hash

    def _get_annotation_metadata(
            self,
            labels: List[LingLabel]
    ) -> Tuple[Dict[str, Dict[str, Any]], Dict[LingLabel, str]]:
        """Store metadata for an annotation."""
        annotation_metadata = {}
        label_to_id = {}
        for label in labels:
            hash_id = self._get_new_hash_id(label, annotation_metadata)
            metadata = {
                'label': label.label,
                'nid': tuple(label.nid),
                'target': label.target
            }
            annotation_metadata[hash_id] = metadata
            label_to_id[label] = hash_id
        return annotation_metadata, label_to_id

    @classmethod
    def from_doc(
            cls,
            document: Union[Path, Document],
            tf_fabric: Fabric,
            project: 'BaseLabelingProject',
            metadata: Dict[str, Any],
    ) -> 'BaseAnnotationSheet':
        """Read in docx annotation sheet."""
        if isinstance(document, Path):
            document = Document(document)

        # populate annotations from docx
        annotations = []
        for table in document.tables:
            for row in table.rows:
                annotations.append(
                    cls._label_from_row(row, metadata)
                )

        # return new class instance
        return cls(
            annotations=annotations,
            tf_fabric=tf_fabric,
            project=project,
            document=document,
        )
